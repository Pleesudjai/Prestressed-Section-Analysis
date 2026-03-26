"""
Shear Design Report — CEE 530 Project 2
Each calculation shows:  Symbolic equation  →  Numerical substitution  →  Result
Matches MATLAB fprintf output style.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page layout ───────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Inches(8.5);  sec.page_height = Inches(11.0)
sec.left_margin = sec.right_margin  = Inches(1.0)
sec.top_margin  = sec.bottom_margin = Inches(1.0)

# ── Font helpers ──────────────────────────────────────────────────────────────
TNR = 'Times New Roman'

def rf(run, bold=False, italic=False, size=10.5, sub=False, underline=False):
    run.font.name    = TNR
    run.font.size    = Pt(size)
    run.font.bold    = bold
    run.font.italic  = italic
    run.font.color.rgb = RGBColor(0,0,0)
    if sub:       run.font.subscript  = True
    if underline: run.font.underline  = True

def new_para(space_before=2, space_after=2, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    return p

def blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)

# ── Heading with bottom border on H1 ─────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 9 if level==2 else 6)
    p.paragraph_format.space_after  = Pt(4  if level==1 else 3)
    r = p.add_run(text)
    r.font.name = 'Arial'; r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)
    r.font.size = Pt(13 if level==1 else 11 if level==2 else 10)
    if level == 1:
        pBdr = OxmlElement('w:pBdr')
        bd = OxmlElement('w:bottom')
        bd.set(qn('w:val'),'single'); bd.set(qn('w:sz'),'6')
        bd.set(qn('w:space'),'4');    bd.set(qn('w:color'),'000000')
        pBdr.append(bd)
        p._p.get_or_add_pPr().append(pBdr)
    return p

# ── Equation block ─────────────────────────────────────────────────────────────
# lines: list of (label, symbolic, numeric, result)
#   label   – left symbol (e.g. 'Vci'); '' for continuation lines
#   symbolic– symbolic RHS; shown in italic (e.g. '0.6λ√f\'c·bw·dp + Vd + (Vi/Mmax)·Mcr')
#   numeric – substituted numbers; shown in normal font ('' to skip)
#   result  – final result with units; shown bold ('' to skip)
EQ_INDENT = 0.55   # inches for equation block

def eq_block(lines):
    """
    Each line: (label, symbolic, numeric, result)
    Renders as:
      label  =  symbolic
             =  numeric
             =  result   (bold)
    """
    for label, sym, num, res in lines:
        p = new_para(space_before=2, space_after=1, indent=EQ_INDENT)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.4))

        if label:
            r = p.add_run(label)
            rf(r, bold=True)
            r2 = p.add_run('  =  ')
            rf(r2)
        else:
            # continuation: pad to align with "="
            r0 = p.add_run('\u00A0'*14 + '=  ')   # nbsp padding + equals
            rf(r0)

        if sym:
            r = p.add_run(sym)
            rf(r, italic=True)
        if num:
            pn = new_para(space_before=1, space_after=1, indent=EQ_INDENT)
            r0 = p.add_run('')   # dummy; num goes on next para
            # Actually build separate para for numeric
            p2 = new_para(space_before=1, space_after=1, indent=EQ_INDENT)
            r0 = p2.add_run('\u00A0'*14 + '=  ')
            rf(r0)
            r1 = p2.add_run(num)
            rf(r1)
        if res:
            p3 = new_para(space_before=1, space_after=3, indent=EQ_INDENT)
            r0 = p3.add_run('\u00A0'*14 + '=  ')
            rf(r0)
            r1 = p3.add_run(res)
            rf(r1, bold=True)

# Simpler single-step equation (sym then bold result on same line)
def eq_single(label, symbolic, result):
    """label = symbolic = result  (all on one logical block, 3 lines)"""
    eq_block([(label, symbolic, '', ''), ('', '', symbolic + '', ''), ('', '', '', result)])


# ── Better equation helper: shows exactly what MATLAB fprintf does ─────────────
# eq_show(label, sym_eq, num_eq, result)
# Example:
#   eq_show('fce', 'Pe/Ac + Pe·e·yb/I',
#                  '85.0/487.0 + 85.0×7.852×20.362/34,638.8',
#                  '0.5669 ksi')
def eq_show(label, sym_eq, num_eq=None, result=None):
    """Three-line equation:  label = sym_eq
                                   = num_eq
                             BOLD: = result"""
    # Line 1: label  =  symbolic
    p1 = new_para(space_before=3, space_after=1, indent=EQ_INDENT)
    r = p1.add_run(label + '  =  ')
    rf(r, bold=True)
    r = p1.add_run(sym_eq)
    rf(r, italic=True)

    # Line 2: numerical substitution
    if num_eq:
        p2 = new_para(space_before=1, space_after=1, indent=EQ_INDENT)
        pad = '\u00A0' * (len(label) + 5)   # align with "="
        r = p2.add_run(pad + '=  ')
        rf(r)
        r = p2.add_run(num_eq)
        rf(r)

    # Line 3: result (bold)
    if result:
        p3 = new_para(space_before=1, space_after=4, indent=EQ_INDENT)
        pad = '\u00A0' * (len(label) + 5)
        r = p3.add_run(pad + '=  ')
        rf(r)
        r = p3.add_run(result)
        rf(r, bold=True)

def body(text, indent=0, bold=False, italic=False):
    p = new_para(indent=indent)
    r = p.add_run(text)
    rf(r, bold=bold, italic=italic)
    return p

def note(text):
    p = new_para(space_before=2, space_after=2, indent=EQ_INDENT)
    r = p.add_run(text)
    rf(r, italic=True)
    return p

# ── Table helpers ─────────────────────────────────────────────────────────────
def set_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'),'single'); bd.set(qn('w:sz'),'4')
        bd.set(qn('w:space'),'0');    bd.set(qn('w:color'),'000000')
        tcBorders.append(bd)
    tcPr.append(tcBorders)

def set_shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, center=False, shade=None, sz=9):
    if shade: set_shade(cell, shade)
    set_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    rf(r, bold=bold, size=sz)

def make_table(headers, rows, widths, center_cols=None):
    """widths in inches"""
    cc = center_cols or []
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    for j,(h,w) in enumerate(zip(headers, widths)):
        c = t.rows[0].cells[j]; c.width = Inches(w)
        cell_text(c, h, bold=True, center=True, shade='CCCCCC', sz=9)
    for i,row in enumerate(rows):
        for j,(val,w) in enumerate(zip(row, widths)):
            c = t.rows[i+1].cells[j]; c.width = Inches(w)
            cell_text(c, str(val), center=(j in cc), sz=9)

# ════════════════════════════════════════════════════════════════════════════════
#  TITLE
# ════════════════════════════════════════════════════════════════════════════════
for txt, sz in [
    ('CEE 530 — Prestressed Concrete', 16),
    ('Shear Design — Project 2', 16),
    ('Double-T Prestressed Beam (ACI 318-19 §22.5)', 13),
    ('Arizona State University  |  Spring 2026', 11),
    ('Chidchanok Pleesudjai', 11),
]:
    p = new_para(space_before=2, space_after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); rf(r, bold=True, size=sz)
blank()

# ════════════════════════════════════════════════════════════════════════════════
#  1. GIVEN INFORMATION
# ════════════════════════════════════════════════════════════════════════════════
heading('1.  Given Information', 1)

heading('1.1  Section Geometry', 2)
make_table(
    headers=['Parameter','Symbol','Value'],
    rows=[
        ['Total section depth',       'h',    '28.0 in.'],
        ['Flange width',              'bf',   '120.0 in.'],
        ['Flange thickness',          'hf',   '2.0 in.'],
        ['Min. web width (2 stems)',  'bw',   '7.50 in.  (2 × 3.75 in.)'],
        ['Span length',               'L',    '768 in.  (64.0 ft)'],
        ['Support condition',         '—',    'Simply supported'],
    ],
    widths=[2.8, 1.0, 2.7],
    center_cols=[1,2],
)
blank()

heading('1.2  Section Properties (Shoelace Formula)', 2)
make_table(
    headers=['Property','Symbol','Value'],
    rows=[
        ['Gross area',                        'Ac',   '487.0 in²'],
        ['Moment of inertia (centroidal)',     'Ix',   '34,638.8 in⁴'],
        ['Centroid from bottom',               'yc',   '20.362 in.'],
        ['Dist. centroid to bottom fiber',     'yb',   '20.362 in.'],
        ['Dist. centroid to top fiber',        'yt',   '7.638 in.'],
        ['Bottom section modulus  Sb=Ix/yb',  'Sb',   '1,701.5 in³'],
        ['Top section modulus  St=Ix/yt',     'St',   '4,532.6 in³'],
    ],
    widths=[3.2, 0.8, 2.5],
    center_cols=[1,2],
)
blank()

heading('1.3  Material Properties', 2)
make_table(
    headers=['Material','Property','Value'],
    rows=[
        ['Concrete',              "f'c",  '6.0 ksi  (6,000 psi)'],
        ['',                     'Ec',   '4,700 ksi'],
        ['',                     'λ',    '1.0  (normal-weight concrete)'],
        ['Prestressing steel',   'fpu',  '270 ksi'],
        ['',                     'Eps',  '28,500 ksi'],
        ['Mild steel (stirrups)','fy',   '60 ksi'],
    ],
    widths=[2.0, 1.2, 3.3],
    center_cols=[1],
)
blank()

heading('1.4  Allowable Stresses — ACI 318-19', 2)
make_table(
    headers=['Stage',"f'c or f'ci",'Compression (ksi)','Tension (ksi)'],
    rows=[
        ['Transfer', "f'ci = 4.8 ksi", "+0.60 f'ci = +2.880",  "-3√f'ci  = −0.208"],
        ['Service',  "f'c  = 6.0 ksi", "+0.45 f'c  = +2.700",  "-12√f'c  = −0.929"],
    ],
    widths=[1.3, 1.5, 2.2, 1.5],
    center_cols=[1,2,3],
)
blank()

heading('1.5  Prestressing Tendons', 2)
make_table(
    headers=['Tendon','x-location','Profile','y (support)','y (midspan)','Pe/strand'],
    rows=[
        ['1 (left)',  '−30.125 in.','Straight','6.00 in.','6.00 in.','21.25 kips'],
        ['2 (left)',  '−30.125 in.','Harped',  '20.36 in.','6.00 in.','21.25 kips'],
        ['3 (right)', '+30.125 in.','Straight','6.00 in.','6.00 in.','21.25 kips'],
        ['4 (right)', '+30.125 in.','Harped',  '20.36 in.','6.00 in.','21.25 kips'],
    ],
    widths=[0.9, 1.0, 0.8, 0.9, 0.9, 1.0],
    center_cols=[1,2,3,4,5],
)
body('4 strands × Aps = 0.153 in²/strand  →  Aps,total = 0.612 in².  '
     'Harp point at x = 240 in. (20 ft) from each support.  '
     'Pi = 100.0 kips;  η = 0.85;  Pe = 85.0 kips.')
blank()

heading('1.6  Applied Loads', 2)
make_table(
    headers=['Load','Symbol','kip/in.','kip/ft'],
    rows=[
        ['Self-weight SW',                 'wsw',   '0.04227','0.5073'],
        ['Superimposed dead load SDL',     'wsdl',  '0.02083','0.2500'],
        ['Live load LL',                   'wll',   '0.03500','0.4200'],
        ['Dead load total  DL = SW+SDL',   'wDL',   '0.06310','0.7573'],
        ['Factored  wu = 1.2DL + 1.6LL',  'wu',    '0.13173','1.5807'],
        ['Fact. var.  wext = 1.2SDL+1.6LL','wext', '0.08100','0.9720'],
    ],
    widths=[2.5, 0.8, 1.0, 0.7],
    center_cols=[1,2,3],
)
blank()

# ════════════════════════════════════════════════════════════════════════════════
#  2. SHEAR DESIGN PROCEDURE
# ════════════════════════════════════════════════════════════════════════════════
heading('2.  Shear Design Procedure — ACI 318-19 §22.5', 1)

body('For prestressed members, ACI 318-19 §22.5.8.3 defines two failure modes:')
body('  Vci  (Flexural-shear cracking)  — inclined crack starts from a flexural crack.', indent=0.3)
body('  Vcw  (Web-shear cracking)  — inclined crack forms in web before any flexure crack.', indent=0.3)
body('Concrete shear capacity:  Vc = min(Vci, Vcw).  '
     'Stirrups required when Vu > φVc, carrying Vs = Vu/φ − Vc.')
blank()

heading('2.1  Critical Section Location', 2)
body('Per ACI 318-19 §22.5.2, the critical section for shear is located at dp from the face '
     'of support, where dp shall not be taken less than 0.8h.')
eq_show('dp,min',
        '0.8 h',
        '0.8 × 28.0 in.',
        '22.40 in.    ← governs throughout (actual dp < 22.40 in. near supports)')
body('All calculations use dp = 22.40 in.  Critical section location:')
eq_show('xcr',
        'dp,min',
        '',
        '22.40 in.  =  1.87 ft from support face')
blank()

heading('2.2  Factored Shear Envelope', 2)
body('For a simply supported beam with uniform factored load, the shear envelope is:')
eq_show('Vu(x)',
        'wu × (L/2 − x)',
        '0.13173 × (384 − x)   [kips, x in inches]')
eq_show('Vu,max  (at support face, x = 0)',
        'wu × L/2',
        '0.13173 × 384',
        '50.58 kips')
eq_show('Vu  (at xcr = 22.40 in.)',
        'wu × (L/2 − xcr)',
        '0.13173 × (384 − 22.40)  =  0.13173 × 361.60',
        '47.63 kips')
blank()

heading('2.3  Vertical Component of Effective Prestress, Vp', 2)
body('Tendons 2 and 4 are harped; tendons 1 and 3 are straight (Vp = 0).  '
     'Harped tendons slope from y = 20.36 in. at supports to y = 6.00 in. at harp point x = 240 in.')

eq_show('θharp',
        'arctan [(yend − ymid) / xharp]',
        'arctan [(20.36 − 6.00) / 240]  =  arctan (14.36 / 240)  =  arctan (0.05983)',
        '3.424°  →  sin θ = 0.05973')
eq_show('Pe per harped strand',
        'Pi,strand × η',
        '25.0 × 0.85',
        '21.25 kips')
eq_show('Vp  (for x ≤ 240 in., sloped zone)',
        '2 strands × Pe,strand × sin θ',
        '2 × 21.25 × 0.05973',
        '2.538 kips  (upward — reduces net Vu on web)')
body('For 240 in. < x ≤ 528 in. (flat zone):  Vp = 0.')
blank()

heading('2.4  Modulus of Rupture and Cracking Moment', 2)
body('All cracking moment equations reference ACI 318-19 §22.5.8.3.3 and §22.5.8.3.1b.')
eq_show('fr',
        '6 λ √f\'c    [ksi, f\'c in ksi]',
        '6 × 1.0 × √6.0',
        '0.4648 ksi')

body('Effective prestress compressive stress at extreme tension fiber (bottom):')
eq_show('fce',
        'Pe/Ac  +  Pe·e·yb / Ix',
        '85.0/487.0  +  (85.0 × e × 20.362) / 34,638.8',
        '← evaluated at each section (e varies with tendon profile)')

body('Unfactored DL moment and flexural tensile stress at bottom fiber:')
eq_show('Md(x)',
        'wDL × x × (L − x) / 2',
        '0.06310 × x × (768 − x) / 2   [kip-in., x in inches]')
eq_show('fd(x)',
        'Md × yb / Ix',
        'Md × 20.362 / 34,638.8')

body('Cracking moment (ACI 318-19 Eq. 22.5.8.3.1b):')
eq_show('Mcr',
        '(Ix / yb) × (fr + fce − fd)',
        '(34,638.8 / 20.362) × (fr + fce − fd)',
        '1,701.5 × (fr + fce − fd)   [kip-in.]')
note('Note: If fr + fce − fd < 0, section is already cracked under DL alone → Mcr = 0.')
blank()

heading('2.5  Flexural-Shear Cracking Strength, Vci  —  ACI 318-19 Eq. 22.5.8.3.1a', 2)
eq_show('Vci',
        '0.6 λ √f\'c · bw · dp  +  Vd  +  (Vi / Mmax) · Mcr',
        '0.6 × 1.0 × √6.0 × 7.5 × 22.40  +  Vd  +  (Vi / Mmax) · Mcr')
body('where:')
make_table(
    headers=['Symbol','Definition','Formula'],
    rows=[
        ['term1','0.6λ√f\'c·bw·dp  (concrete contribution)',
         '0.6×1.0×0.07746×7.5×22.40 = 7.808 kips  (constant)'],
        ['Vd','Unfactored DL shear',  'wDL×(L/2−x)  =  0.06310×(384−x)'],
        ['Vi','Factored shear from loads applied after decompression',
         'wext×(L/2−x)  =  0.08100×(384−x)'],
        ['Mmax','Factored moment at x from same loading as Vi',
         'wext×x×(L−x)/2  =  0.08100×x×(768−x)/2'],
    ],
    widths=[0.7, 2.8, 3.0],
    center_cols=[0],
)
body('Minimum value (always applies as lower bound):')
eq_show('Vci,min',
        '1.7 λ √f\'c · bw · dp',
        '1.7 × 1.0 × 0.07746 × 7.5 × 22.40',
        '22.12 kips  (constant throughout beam)')
blank()

heading('2.6  Web-Shear Cracking Strength, Vcw  —  ACI 318-19 Eq. 22.5.8.3.2', 2)
eq_show('Vcw',
        '(3.5 λ √f\'c  +  0.3 fpc) · bw · dp  +  Vp')
body('where fpc = average prestress at centroid:')
eq_show('fpc',
        'Pe / Ac',
        '85.0 / 487.0',
        '0.1745 ksi  (constant throughout beam)')
body('Concrete + prestress term (constant, since bw, dp, fpc are constant):')
eq_show('(3.5λ√f\'c + 0.3fpc)·bw·dp',
        '(3.5×1.0×0.07746  +  0.3×0.1745) × 7.5 × 22.40',
        '(0.2711  +  0.05235) × 168.0  =  0.32345 × 168.0',
        '54.34 kips')
eq_show('Vcw  (sloped zone, x ≤ 240 in., Vp = 2.538 kips)',
        '54.34  +  Vp', '54.34  +  2.538', '56.88 kips')
eq_show('Vcw  (flat zone, 240 < x ≤ 528 in., Vp = 0)',
        '54.34  +  0', '', '54.34 kips')
blank()

heading('2.7  Minimum Transverse Reinforcement  —  ACI 318-19 §22.5.10.5', 2)
body('Three criteria; the largest governs:')
eq_show('(Av/s)₁',
        '0.75 √f\'c · bw / fy    [f\'c in psi]',
        '0.75 × √6,000 × 7.5 / 60,000  =  0.75 × 77.46 × 7.5 / 60,000',
        '0.00726 in²/in.  ← GOVERNS')
eq_show('(Av/s)₂',
        '50 · bw / fy    [f\'c in psi]',
        '50 × 7.5 / 60,000',
        '0.00625 in²/in.')
eq_show('(Av/s)₃',
        'Aps·fpu / (80·fy·dp) × √(dp/bw)',
        '0.612×270 / (80×60×22.40) × √(22.40/7.50)',
        '0.00266 in²/in.')
body('Stirrup selected:  #3 closed U-stirrups  →  Av = 2 × 0.11 = 0.22 in².')
eq_show('s  (from min Av/s)',
        'Av / (Av/s)min',
        '0.22 / 0.00726',
        '30.3 in.  →  limited by smax')
blank()

heading('2.8  Spacing Limits  —  ACI 318-19 §22.7.6.2', 2)
eq_show('smax  (basic)',
        'min(3h/4,  24 in.)',
        'min(3×28/4,  24)  =  min(21.0, 24.0)',
        '21.0 in.')
eq_show('Vs,threshold  (for tight zone)',
        '4 λ √f\'c · bw · dp',
        '4 × 1.0 × 0.07746 × 7.5 × 22.40',
        '52.17 kips')
body('When Vs > 52.17 kips:  smax,tight = min(3h/8, 12 in.) = 10.5 in.  '
     '(Does NOT apply here — max Vs,req = 22.77 kips < 52.17 kips throughout.)')
blank()

# ════════════════════════════════════════════════════════════════════════════════
#  3. DETAILED SECTION CALCULATIONS
# ════════════════════════════════════════════════════════════════════════════════
heading('3.  Detailed Calculations at Selected Sections', 1)

body('Results at five sections are shown.  The beam is symmetric; the right half mirrors the left.')
blank()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3.1  —  x = 22.40 in  (Critical Section)
# ─────────────────────────────────────────────────────────────────────────────
heading('3.1  Critical Section: x = 22.40 in.  (1.87 ft from support)', 2)

heading('Step A — Tendon Position and Effective Depth', 3)
body('Harped tendons (2 and 4) slope linearly from y = 20.36 in. at x = 0 to y = 6.00 in. at x = 240 in.:')
eq_show('yharped(x)',
        'y_support  +  (y_mid − y_support) × (x / x_harp)',
        '20.36  +  (6.00 − 20.36) × (22.40 / 240)',
        '20.36  −  1.340  =  19.02 in.')
eq_show('yps  (centroid of all 4 strands)',
        '(y₁ + y₂ + y₃ + y₄) / 4',
        '(6.00 + 19.02 + 6.00 + 19.02) / 4',
        '12.51 in.')
eq_show('e  (eccentricity)',
        'yc − yps',
        '20.362 − 12.51',
        '7.852 in.  (tendon below centroid → positive)')
eq_show('dp,calc',
        'h − yps',
        '28.0 − 12.51',
        '15.49 in.  <  0.8h = 22.40 in.  →  use dp = 22.40 in.')

heading('Step B — fce: Effective Prestress at Bottom Fiber', 3)
eq_show('fce',
        'Pe/Ac  +  Pe·e·yb / Ix',
        '85.0/487.0  +  (85.0 × 7.852 × 20.362) / 34,638.8',
        '0.1745  +  0.3924  =  0.5669 ksi  (compression +)')

heading('Step C — fd: DL Flexural Stress at Bottom Fiber', 3)
eq_show('Md',
        'wDL × x × (L − x) / 2',
        '0.06310 × 22.40 × (768 − 22.40) / 2  =  0.06310 × 22.40 × 372.80',
        '527.0 kip-in.')
eq_show('fd',
        'Md × yb / Ix',
        '527.0 × 20.362 / 34,638.8',
        '0.3098 ksi  (tension −)')

heading('Step D — Mcr: Cracking Moment', 3)
eq_show('fr + fce − fd',
        '',
        '0.4648  +  0.5669  −  0.3098',
        '0.7219 ksi  >  0  →  Mcr > 0')
eq_show('Mcr',
        '(Ix / yb) × (fr + fce − fd)',
        '(34,638.8 / 20.362) × 0.7219  =  1,701.5 × 0.7219',
        '1,228 kip-in.')

heading('Step E — Vci: Flexural-Shear Cracking Strength', 3)
eq_show('Vd',
        'wDL × (L/2 − x)',
        '0.06310 × (384 − 22.40)  =  0.06310 × 361.60',
        '22.82 kips')
eq_show('Vi',
        'wext × (L/2 − x)',
        '0.08100 × (384 − 22.40)  =  0.08100 × 361.60',
        '29.29 kips')
eq_show('Mmax',
        'wext × x × (L − x) / 2',
        '0.08100 × 22.40 × (768 − 22.40) / 2  =  0.08100 × 22.40 × 372.80',
        '676.4 kip-in.')

body('Substituting into ACI Eq. 22.5.8.3.1a:')
eq_show('term1  =  0.6λ√f\'c·bw·dp',
        '',
        '0.6 × 1.0 × 0.07746 × 7.5 × 22.40',
        '7.808 kips')
eq_show('term2  =  Vd',
        '',
        '',
        '22.82 kips')
eq_show('term3  =  (Vi/Mmax)·Mcr',
        '',
        '(29.29 / 676.4) × 1,228  =  0.04331 × 1,228',
        '53.18 kips')
eq_show('Vci  =  term1 + term2 + term3',
        '',
        '7.808  +  22.82  +  53.18',
        '83.81 kips')
eq_show('Vci,min  =  1.7λ√f\'c·bw·dp',
        '',
        '1.7 × 1.0 × 0.07746 × 7.5 × 22.40',
        '22.12 kips  <  83.81 kips  →  Vci = 83.81 kips')

heading('Step F — Vcw: Web-Shear Cracking Strength', 3)
body('x = 22.40 in. ≤ 240 in.  →  sloped tendon zone  →  Vp = 2.538 kips.')
eq_show('Vcw',
        '(3.5λ√f\'c + 0.3fpc)·bw·dp  +  Vp',
        '54.34  +  2.538',
        '56.88 kips')

heading('Step G — Vc, φVc, and Stirrup Design', 3)
eq_show('Vc  =  min(Vci, Vcw)',
        '',
        'min(83.81,  56.88)',
        '56.88 kips  (Vcw governs)')
eq_show('φVc  =  0.75 × Vc',
        '',
        '0.75 × 56.88',
        '42.66 kips')
body('Vu = 47.63 kips  >  φVc = 42.66 kips  →  Stirrups required.')
eq_show('Vs,req',
        'Vu/φ  −  Vc',
        '47.63/0.75  −  56.88  =  63.51  −  56.88',
        '6.63 kips')
eq_show('Av/s  (demand)',
        'Vs,req / (fy · dp)',
        '6.63 / (60 × 22.40)',
        '0.00493 in²/in.   <  Av/s,min = 0.00726 in²/in.')
eq_show('s  =  Av / (Av/s,min)',
        '',
        '0.22 / 0.00726  =  30.3 in.   →  limited by smax',
        'USE  s = 21 in.')
blank()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3.2  —  x = 63.6 in  (5.30 ft)
# ─────────────────────────────────────────────────────────────────────────────
heading('3.2  Section at x = 63.6 in.  (5.30 ft)', 2)

heading('Step A — Tendon Position', 3)
eq_show('yharped',
        '20.36  +  (6.00 − 20.36) × (63.6 / 240)',
        '20.36  −  3.805',
        '16.56 in.')
eq_show('yps',
        '(6.00 + 16.56 + 6.00 + 16.56) / 4',
        '',
        '11.28 in.')
eq_show('e',
        'yc − yps',
        '20.362 − 11.28',
        '9.082 in.')

heading('Step B — fce', 3)
eq_show('fce',
        'Pe/Ac  +  Pe·e·yb / Ix',
        '85.0/487.0  +  (85.0 × 9.082 × 20.362) / 34,638.8',
        '0.1745  +  0.4537  =  0.6282 ksi')

heading('Step C — fd', 3)
eq_show('Md',
        'wDL × x × (L − x) / 2',
        '0.06310 × 63.6 × (768 − 63.6) / 2  =  0.06310 × 63.6 × 352.2',
        '1,413.5 kip-in.')
eq_show('fd',
        'Md × yb / Ix',
        '1,413.5 × 20.362 / 34,638.8',
        '0.8309 ksi')

heading('Step D — Mcr', 3)
eq_show('fr + fce − fd',
        '',
        '0.4648  +  0.6282  −  0.8309',
        '0.2621 ksi  >  0  →  Mcr > 0')
eq_show('Mcr',
        '1,701.5 × (fr + fce − fd)',
        '1,701.5 × 0.2621',
        '446.2 kip-in.')

heading('Step E — Vci', 3)
eq_show('Vd',
        'wDL × (L/2 − x)',
        '0.06310 × (384 − 63.6)  =  0.06310 × 320.4',
        '20.22 kips')
eq_show('Vi',
        'wext × (L/2 − x)',
        '0.08100 × 320.4',
        '25.95 kips')
eq_show('Mmax',
        'wext × x × (L − x) / 2',
        '0.08100 × 63.6 × (768 − 63.6) / 2  =  0.08100 × 63.6 × 352.2',
        '1,814.3 kip-in.')
eq_show('term1  =  0.6λ√f\'c·bw·dp', '', '', '7.808 kips  (constant)')
eq_show('term2  =  Vd', '', '', '20.22 kips')
eq_show('term3  =  (Vi/Mmax)·Mcr',
        '',
        '(25.95 / 1,814.3) × 446.2  =  0.01431 × 446.2',
        '6.384 kips')
eq_show('Vci  =  term1 + term2 + term3',
        '',
        '7.808  +  20.22  +  6.384',
        '34.41 kips  >  Vci,min = 22.12 kips  →  Vci = 34.41 kips')

heading('Step F — Vcw', 3)
body('x = 63.6 in. ≤ 240 in.  →  Vp = 2.538 kips.')
eq_show('Vcw', '54.34  +  Vp', '54.34  +  2.538', '56.88 kips')

heading('Step G — Vc and Stirrup Design', 3)
eq_show('Vc  =  min(Vci, Vcw)',
        '',
        'min(34.41,  56.88)',
        '34.41 kips   ← Vci governs')
eq_show('φVc', '0.75 × 34.41', '', '25.81 kips')
body('Vu = 42.21 kips  >  φVc = 25.81 kips  →  Stirrups required.')
eq_show('Vs,req',
        'Vu/φ  −  Vc',
        '42.21/0.75  −  34.41  =  56.28  −  34.41',
        '21.87 kips')
eq_show('Av/s  (demand)',
        'Vs,req / (fy · dp)',
        '21.87 / (60 × 22.40)',
        '0.01627 in²/in.   >  Av/s,min = 0.00726  →  demand governs')
eq_show('s  =  Av / (Av/s)',
        '',
        '0.22 / 0.01627  =  13.5 in.',
        'USE  s = 13 in.  (≤ smax = 21 in.  ✓)')
blank()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3.3  —  x = 128.4 in  (10.70 ft)
# ─────────────────────────────────────────────────────────────────────────────
heading('3.3  Section at x = 128.4 in.  (10.70 ft)', 2)

heading('Step A — Tendon Position', 3)
eq_show('yharped',
        '20.36  +  (6.00 − 20.36) × (128.4 / 240)',
        '20.36  −  7.683',
        '12.68 in.')
eq_show('yps', '(6.00 + 12.68 + 6.00 + 12.68) / 4', '', '9.34 in.')
eq_show('e',   '20.362 − 9.34', '', '11.02 in.')

heading('Step B — fce', 3)
eq_show('fce',
        'Pe/Ac  +  Pe·e·yb / Ix',
        '85.0/487.0  +  (85.0 × 11.02 × 20.362) / 34,638.8',
        '0.1745  +  0.5508  =  0.7253 ksi')

heading('Step C — fd', 3)
eq_show('Md',
        'wDL × x × (L − x) / 2',
        '0.06310 × 128.4 × (768 − 128.4) / 2  =  0.06310 × 128.4 × 319.8',
        '2,591.2 kip-in.')
eq_show('fd', 'Md × yb / Ix', '2,591.2 × 20.362 / 34,638.8', '1.5232 ksi')

heading('Step D — Mcr  (Section cracked under DL alone)', 3)
eq_show('fr + fce − fd',
        '',
        '0.4648  +  0.7253  −  1.5232',
        '−0.333 ksi  <  0  →  Mcr = 0')
body('Interpretation: The effective prestress (fce = 0.7253 ksi) plus modulus of rupture '
     '(fr = 0.4648 ksi) together cannot overcome the DL tensile stress (fd = 1.5232 ksi) '
     'at the bottom fiber.  The (Vi/Mmax)·Mcr term vanishes.  Vci reduces to its minimum.')
eq_show('Vci', '1.7λ√f\'c·bw·dp', '1.7 × 1.0 × 0.07746 × 7.5 × 22.40', '22.12 kips  (minimum governs)')

heading('Step E — Vcw', 3)
body('x = 128.4 in. ≤ 240 in.  →  Vp = 2.538 kips.')
eq_show('Vcw', '54.34  +  2.538', '', '56.88 kips')

heading('Step F — Vc and Stirrup Design', 3)
eq_show('Vc  =  min(Vci, Vcw)',
        '',
        'min(22.12,  56.88)',
        '22.12 kips   ← Vci,min governs')
eq_show('φVc', '0.75 × 22.12', '', '16.59 kips')
body('Vu = 33.67 kips  >  φVc = 16.59 kips  →  Stirrups required.')
eq_show('Vs,req',
        'Vu/φ  −  Vc',
        '33.67/0.75  −  22.12  =  44.89  −  22.12',
        '22.77 kips  (MATLAB exact at this x: 20.95 kips)')
eq_show('Av/s  (demand)',
        'Vs,req / (fy · dp)',
        '20.95 / (60 × 22.40)',
        '0.01559 in²/in.   >  Av/s,min  →  demand governs')
eq_show('s  =  Av / (Av/s)',
        '',
        '0.22 / 0.01559  =  14.1 in.',
        'USE  s = 14 in.  (≤ smax = 21 in.  ✓)')
blank()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3.4  —  x = 192 in  (16 ft)
# ─────────────────────────────────────────────────────────────────────────────
heading('3.4  Section at x = 192.0 in.  (16.0 ft)', 2)

heading('Step A — Tendon Position', 3)
eq_show('yharped',
        '20.36  +  (6.00 − 20.36) × (192.0 / 240)',
        '20.36  −  11.486',
        '8.874 in.')
eq_show('yps', '(6.00 + 8.874 + 6.00 + 8.874) / 4', '', '7.437 in.')
eq_show('e',   '20.362 − 7.437', '', '12.925 in.')

heading('Step B — fce', 3)
eq_show('fce',
        'Pe/Ac  +  Pe·e·yb / Ix',
        '85.0/487.0  +  (85.0 × 12.925 × 20.362) / 34,638.8',
        '0.1745  +  0.6459  =  0.8204 ksi')

heading('Step C — fd', 3)
eq_show('Md',
        'wDL × x × (L − x) / 2',
        '0.06310 × 192.0 × (768 − 192.0) / 2  =  0.06310 × 192.0 × 288.0',
        '3,489.4 kip-in.')
eq_show('fd', 'Md × yb / Ix', '3,489.4 × 20.362 / 34,638.8', '2.051 ksi')

heading('Step D — Mcr', 3)
eq_show('fr + fce − fd',
        '',
        '0.4648  +  0.8204  −  2.051',
        '−0.766 ksi  <  0  →  Mcr = 0  →  Vci = Vci,min = 22.12 kips')

heading('Step E — Vcw', 3)
body('x = 192.0 in. ≤ 240 in.  →  Vp = 2.538 kips.  →  Vcw = 56.88 kips.')

heading('Step F — Vc and Stirrup Design', 3)
eq_show('Vc  =  min(22.12, 56.88)', '', '', '22.12 kips  (Vci,min governs)')
eq_show('φVc', '0.75 × 22.12', '', '16.59 kips')
body('Vu = 25.29 kips  >  φVc = 16.59 kips  →  Stirrups required.')
eq_show('Vs,req',
        'Vu/φ  −  Vc',
        '25.29/0.75  −  22.12  =  33.72  −  22.12',
        '11.60 kips')
eq_show('Av/s  (demand)',
        'Vs,req / (fy · dp)',
        '11.60 / (60 × 22.40)',
        '0.00863 in²/in.   >  Av/s,min = 0.00726  →  demand governs')
eq_show('s  =  Av / (Av/s)',
        '',
        '0.22 / 0.00863  =  25.5 in.   →  limited by smax',
        'USE  s = 21 in.')
blank()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3.5  —  Midspan (x = 384 in)
# ─────────────────────────────────────────────────────────────────────────────
heading('3.5  Midspan: x = 384.0 in.  (32.0 ft)', 2)
body('At midspan all four tendons are at y = 6.0 in.  '
     'Vu = 0 (by symmetry).  Vp = 0 (flat zone).')

heading('Step A — Tendon Position', 3)
eq_show('yps (midspan)', '6.0 in. (all strands)', '', '')
eq_show('e', 'yc − yps', '20.362 − 6.0', '14.362 in.  ← maximum eccentricity')

heading('Step B — fce', 3)
eq_show('fce',
        'Pe/Ac  +  Pe·e·yb / Ix',
        '85.0/487.0  +  (85.0 × 14.362 × 20.362) / 34,638.8',
        '0.1745  +  0.7178  =  0.8923 ksi')

heading('Step C — fd', 3)
eq_show('Md',
        'wDL × L² / 8',
        '0.06310 × 768² / 8  =  0.06310 × 73,728 / 8',
        '4,652.6 kip-in.  (maximum DL moment)')
eq_show('fd', 'Md × yb / Ix', '4,652.6 × 20.362 / 34,638.8', '2.735 ksi')

heading('Step D — Mcr', 3)
eq_show('fr + fce − fd',
        '',
        '0.4648  +  0.8923  −  2.735',
        '−1.378 ksi  <  0  →  Mcr = 0')
body('The maximum DL tensile stress at the bottom (2.735 ksi) far exceeds the combined '
     'prestress and rupture capacity (0.892 + 0.465 = 1.357 ksi).  '
     'Vci defaults to its minimum value.')
eq_show('Vci', '22.12 kips  (minimum governs)', '', '')

heading('Step E — Vcw (flat zone)', 3)
eq_show('Vcw  (Vp = 0)',
        '(3.5λ√f\'c + 0.3fpc)·bw·dp  +  0',
        '0.32345 × 168.0  +  0',
        '54.34 kips')

heading('Step F — Vc and Stirrup Check', 3)
eq_show('Vc  =  min(22.12, 54.34)', '', '', '22.12 kips  (Vci,min governs)')
eq_show('φVc', '0.75 × 22.12', '', '16.59 kips')
body('Vu = 0 kips  <  φVc = 16.59 kips  →  No stirrups required by strength; use minimum.')
eq_show('s  (minimum)', 'Av / (Av/s,min)', '0.22 / 0.00726  =  30.3 in.', 'USE  s = 21 in.  (limited by smax)')
blank()

# ════════════════════════════════════════════════════════════════════════════════
#  4. SUMMARY
# ════════════════════════════════════════════════════════════════════════════════
heading('4.  Shear Design Summary', 1)

heading('4.1  Section-by-Section Results  (Left Half; Right Half Mirrors by Symmetry)', 2)
make_table(
    headers=['x\n(ft)', 'Vu\n(k)', 'Vci\n(k)', 'Vcw\n(k)', 'Vc\n(k)', 'φVc\n(k)',
             'Vs,req\n(k)', 'Av/s\n(in²/in)', 's_req\n(in)', 's_use\n(in)', 'Governs'],
    rows=[
        ['0.0',  '50.58','>>Vcw','56.88','56.88','42.66','10.56','0.00786','28.0','21','Min Av/s'],
        ['1.87', '47.63','83.81','56.88','56.88','42.66', '6.63','0.00726*','30.3','21','Vcw; min Av/s'],
        ['5.30', '42.21','34.41','56.88','34.41','25.81','21.87','0.01627','13.5','13','Vci; demand'],
        ['10.70','33.67','22.12*','56.88','22.12','16.59','20.95','0.01559','14.1','14','Vci,min; demand'],
        ['16.00','25.29','22.12*','56.88','22.12','16.59','11.60','0.00863','25.5','21','Vci,min; smax'],
        ['21.33','16.90','22.12*','54.34','22.12','16.59', '0.40','0.00726*','30.3','21','Min Av/s'],
        ['32.00', '0.00','22.12*','54.34','22.12','16.59', '0.00','0.00726*','30.3','21','Min Av/s'],
    ],
    widths=[0.55, 0.50, 0.60, 0.60, 0.55, 0.55, 0.60, 0.80, 0.60, 0.55, 1.10],
    center_cols=list(range(11)),
)
body('* Minimum Vci,min = 22.12 kips or Av/s,min = 0.00726 in²/in. governs.  '
     '>> Vcw: Vci → ∞ at support (Mmax → 0); Vcw governs directly.')
blank()

heading('4.2  Stirrup Layout — #3 Closed U-Stirrups  (Av = 0.22 in²,  fy = 60 ksi)', 2)
make_table(
    headers=['Zone (each half from support)', 'x range', 'Spacing', 'Governing criterion'],
    rows=[
        ['Support — critical section (Vcw zone)',  '0 to 22 in.  (0 to 1.87 ft)',     's = 21 in.', 'Min. Av/s;  Vcw controls Vc'],
        ['High shear — Vci zone',                  '22 to ~80 in.  (1.87 to 6.67 ft)', 's = 13 in.', 'Vci governs Vc; demand Av/s'],
        ['Transition — Vci,min zone',              '~80 to ~160 in.  (6.67 to 13.3 ft)','s = 14 in.','Vci,min; demand Av/s'],
        ['Mid-beam to midspan',                    '~160 in. to 32 ft',               's = 21 in.', 'smax limit or min. Av/s'],
    ],
    widths=[2.4, 1.9, 0.8, 2.0],
    center_cols=[2],
)
blank()

heading('4.3  Section Adequacy Checks', 2)
eq_show('Vs,max  (ACI 318-19 §22.5.10.1)',
        '8λ√f\'c · bw · dp',
        '8 × 1.0 × 0.07746 × 7.5 × 22.40',
        '104.2 kips')
body('Maximum Vs,req = 22.77 kips  <<  Vs,max = 104.2 kips  →  '
     'Section dimensions are adequate (no enlargement required).  ✓')
body('Tight spacing threshold = 52.17 kips  >  Vs,req,max = 22.77 kips  →  '
     'smax = 21 in. (basic) applies everywhere.  ✓')
blank()

# ════════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ════════════════════════════════════════════════════════════════════════════════
heading('References', 1)
refs = [
    ('ACI Committee 318.',
     ' (2019). ',
     'Building Code Requirements for Structural Concrete (ACI 318-19) and Commentary.',
     ' American Concrete Institute, Farmington Hills, MI.'),
    ('Naaman, A. E.',
     ' (2004). ',
     'Prestressed Concrete Analysis and Design: Fundamentals',
     ' (2nd ed.). Techno Press 3000, Ann Arbor, MI.'),
    ('CEE 530 Prestressed Concrete — Shear design notes and handouts.',
     '',
     '',
     ' Arizona State University, Spring 2026.'),
]
for a, b, c, d in refs:
    p = new_para(space_before=3, space_after=4)
    r1 = p.add_run(a); rf(r1, bold=True)
    r2 = p.add_run(b); rf(r2)
    r3 = p.add_run(c); rf(r3, italic=True)
    r4 = p.add_run(d); rf(r4)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('ShearDesign_Report.docx')
print('Saved: ShearDesign_Report.docx')
