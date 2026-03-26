"""
CEE 530 Prestressed Concrete — Exam Preparation Document
Generated with python-docx (guaranteed Word-compatible)
All parameters from Project 1 (Double-T beam, ASU Spring 2026)
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r'C:\Users\chidc\ASU Dropbox\Chidchanok Pleesudjai\PhD COURSES\2026 Spring\CEE 530 Prestressed Concrete\Prestressed-Section-Analysis\06_draft_text\Exam_Prep_CEE530.docx'

doc = Document()

# ── Page setup: US Letter, 0.75-in margins ─────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(0.85)
section.right_margin  = Inches(0.85)
section.top_margin    = Inches(0.85)
section.bottom_margin = Inches(0.85)

# ── Style helpers ──────────────────────────────────────────────────────────
def set_run(run, bold=False, italic=False, size=11, mono=False, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Courier New' if mono else 'Arial'
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_run(para, text, bold=False, italic=False, size=11, mono=False, sub=False, sup=False):
    run = para.add_run(text)
    set_run(run, bold=bold, italic=italic, size=size, mono=mono)
    if sub:
        run.font.subscript = True
    if sup:
        run.font.superscript = True
    return run

def heading1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(5)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text='', align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if text:
        add_run(p, text)
    return p

def formula(lines, shaded=True):
    """Add a shaded formula box (single-column table)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    # Shade the cell
    if shaded:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        tc_pr.append(shd)
    # Clear default paragraph
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if isinstance(line, str):
            add_run(p, line, mono=True, size=10)
        elif isinstance(line, list):
            # list of (text, kwargs) tuples
            for item in line:
                if isinstance(item, str):
                    add_run(p, item, size=10)
                else:
                    text, kw = item
                    add_run(p, text, size=10, **kw)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)
    return tbl

def warn(lines, fill='EBEBEB'):
    """Add a warning/pitfall box with left border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)
    # Thick left border
    tc_bdr = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'thick' if side == 'left' else 'single')
        b.set(qn('w:sz'), '18' if side == 'left' else '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
        tc_bdr.append(b)
    tc_pr.append(tc_bdr)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if isinstance(line, str):
            add_run(p, line, size=10)
        elif isinstance(line, list):
            for item in line:
                if isinstance(item, str):
                    add_run(p, item, size=10)
                else:
                    text, kw = item
                    add_run(p, text, size=10, **kw)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)
    return tbl

def data_table(headers, rows, col_widths_in):
    """Add a data table. col_widths_in is list of column widths in inches."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Inches(col_widths_in[i])
        # Gray shading
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'DDDDDD')
        tc_pr.append(shd)
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        if isinstance(h, list):
            for item in h:
                if isinstance(item, str): add_run(p, item, bold=True, size=9.5)
                else:
                    text, kw = item; add_run(p, text, bold=True, size=9.5, **kw)
        else:
            add_run(p, h, bold=True, size=9.5)
    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri+1]
        for ci, cell_data in enumerate(row):
            cell = tr.cells[ci]
            cell.width = Inches(col_widths_in[ci])
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            if isinstance(cell_data, list):
                for item in cell_data:
                    if isinstance(item, str): add_run(p, item, size=9.5)
                    else:
                        text, kw = item; add_run(p, text, size=9.5, **kw)
            else:
                add_run(p, str(cell_data), size=9.5)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)
    return tbl

def two_col(pairs, w1=2.5, total=6.9):
    """Label | Value two-column table."""
    w2 = total - w1
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.style = 'Table Grid'
    for ri, (label, value) in enumerate(pairs):
        # Label cell
        lc = tbl.rows[ri].cells[0]
        lc.width = Inches(w1)
        tc_pr = lc._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F0F0F0')
        tc_pr.append(shd)
        lc.paragraphs[0]._p.getparent().remove(lc.paragraphs[0]._p)
        lp = lc.add_paragraph()
        lp.paragraph_format.space_before = Pt(1); lp.paragraph_format.space_after = Pt(1)
        if isinstance(label, list):
            for item in label:
                if isinstance(item, str): add_run(lp, item, bold=True, size=9.5)
                else:
                    text, kw = item; add_run(lp, text, bold=True, size=9.5, **kw)
        else:
            add_run(lp, label, bold=True, size=9.5)
        # Value cell
        vc = tbl.rows[ri].cells[1]
        vc.width = Inches(w2)
        vc.paragraphs[0]._p.getparent().remove(vc.paragraphs[0]._p)
        vp = vc.add_paragraph()
        vp.paragraph_format.space_before = Pt(1); vp.paragraph_format.space_after = Pt(1)
        if isinstance(value, list):
            for item in value:
                if isinstance(item, str): add_run(vp, item, mono=True, size=9.5)
                else:
                    text, kw = item; add_run(vp, text, size=9.5, **kw)
        else:
            add_run(vp, str(value), mono=True, size=9.5)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)
    return tbl

sp = lambda: doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════
sp(); sp()
p = body('CEE 530  -  Prestressed Concrete', align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p, '', bold=True, size=20)   # clear default
p.clear()
add_run(p, 'CEE 530  -  Prestressed Concrete', bold=True, size=20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = body('', align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p2, 'Exam Preparation Notes', bold=True, size=16)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = body('', align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p3, 'Open-Book Reference  |  Hand Calculation Procedures  |  Project 1 Parameters', size=10)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

sp()
p4 = body('', align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p4, 'Student: Chidchanok Pleesudjai (Fen)', size=10)
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p5 = body('', align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p5, 'Arizona State University  |  Spring 2026', size=10)
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
p6 = body('', align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p6, 'Instructor: Prof. Barzin Mobasher', size=10)
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

sp()
warn([
    [('SIGN CONVENTION: ', {'bold': True}), 'Compression = Positive (+)   |   Tension = Negative (-)'],
    [('Origin: ', {'bold': True}), 'y = 0 at bottom of stems, upward positive. Eccentricity e = y_c - y_tendon (positive when tendon is BELOW centroid)'],
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 1 — SECTION PROPERTIES
# ══════════════════════════════════════════════════════════════════════════
heading1('1   Section Properties')

heading2('1.1  Shoelace Formula  (Polygon Vertices, CCW Order)')
body('List vertices (x_i, y_i) counterclockwise. Close polygon by repeating vertex 1 at the end.')
sp()
formula([
    '  Area:',
    '    A = (1/2) * |SUM(x_i * y_{i+1}  -  x_{i+1} * y_i)|',
    '',
    '  Centroid (y from origin):',
    '    y_c = SUM[(y_i + y_{i+1}) * (x_i*y_{i+1} - x_{i+1}*y_i)] / (6A)',
    '',
    '  Centroidal moment of inertia:',
    '    I_c = SUM[(y_i^2 + y_i*y_{i+1} + y_{i+1}^2) * (x_i*y_{i+1} - x_{i+1}*y_i)] / 12',
])

heading2('1.2  Section Moduli and Kern Points')
formula([
    '  Distances from centroid to extreme fibers:',
    '    y_t = y_max - y_c          (centroid to top fiber)',
    '    y_b = y_c - y_min          (centroid to bottom fiber = y_c when y_min = 0)',
    '',
    '  Section moduli:',
    '    S_t = I_c / y_t            (top fiber modulus, in^3)',
    '    S_b = I_c / y_b            (bottom fiber modulus, in^3)',
    '',
    '  Kern points (no-tension zone boundaries):',
    '    k_t = I_c / (A_c * y_t) = r^2 / y_t    (kern toward top, in)',
    '    k_b = I_c / (A_c * y_b) = r^2 / y_b    (kern toward bottom, in)',
    '    r^2 = I_c / A_c                          (radius of gyration squared)',
])
sp()
warn([
    [('Physical meaning of kern points: ', {'bold': True}),
     'If the prestress resultant falls within the kern (k_b below centroid to k_t above), '
     'the entire cross-section stays in compression. Tendon eccentricity typically exceeds k_b for efficient design.'],
])

heading2('1.3  Project 1  (Double-T, y = 0 at stem soffit)')
two_col([
    ('Vertex coordinates (CCW):', '(-60,26),(-33,26),(-32,0),(-28.25,0),(-27.25,26),(27.25,26),(28.25,0),(32,0),(33,26),(60,26),(60,28),(-60,28)'),
    ('Area  Ac', '487.0 in^2'),
    ('Moment of inertia  Ic', '34,638.8 in^4'),
    ('Centroid  yc (from bottom)', '20.362 in'),
    ('Top fiber distance  yt', '28.0 - 20.362 = 7.638 in'),
    ('Bottom fiber distance  yb', '20.362 in'),
    ('Top modulus  St', '34638.8 / 7.638 = 4,535.1 in^3'),
    ('Bottom modulus  Sb', '34638.8 / 20.362 = 1,701.1 in^3'),
    ('Kern  kt', '9.312 in'),
    ('Kern  kb', '3.493 in'),
    ('Radius of gyration  r^2 = Ic/Ac', '71.13 in^2'),
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 2 — STRESS ANALYSIS
# ══════════════════════════════════════════════════════════════════════════
heading1('2   Stress Analysis Along the Beam')

heading2('2.1  Loads and Moments')
body('For a simply supported beam (moment at position x from left support):')
sp()
formula([
    '  M(x) = w * x * (L - x) / 2       (kip-in,  x and L in inches, w in kip/in)',
    '',
    '  M_max = w * L^2 / 8               (at midspan, x = L/2)',
])
sp()
two_col([
    ('Transfer moment  Mi', 'Mi = w_sw * x * (L-x) / 2    [self-weight only]'),
    ('Service moment  MT', 'MT = (w_sw + w_SDL + w_LL) * x * (L-x) / 2'),
    ('--- Project 1 Loads ---', ''),
    ('  Self-weight  w_sw', 'Ac * (150/1728) / 1000 = 0.04227 kip/in'),
    ('  SDL (2-in topping) w_SDL', '2 x 120 x (150/1728) / 1000 = 0.02083 kip/in'),
    ('  Live load  w_LL', '420 lb/ft / 12 / 1000 = 0.03500 kip/in'),
    ('  Total service  w_s', '0.09810 kip/in'),
    ('Span  L', '(20+24+20) x 12 = 768 in  (64 ft)'),
])

heading2('2.2  Eccentricity of the Tendon Group  e_cg')
body('When multiple tendons are at different heights, combine into one resultant F at eccentricity e_cg:')
sp()
formula([
    '  Force-weighted centroid eccentricity:',
    '    e_cg(x) = SUM[ Pi * e_i(x) ] / SUM[ Pi ]     where  Pi = Aps_i * fpi_i',
    '',
    '  Individual tendon eccentricity:',
    '    e_i(x) = y_c - y_i(x)     (positive when tendon is BELOW centroid)',
    '',
    '  Harped tendon y-profile (trapezoidal):',
    '    y_h(x) = y_sup + (y_drape - y_sup) * x / x_drape     for 0 <= x <= x_drape',
    '    y_h(x) = y_drape                                       for x > x_drape',
    '',
    '  NOTE: Use FORCE-weighted average (Pi = Aps*fpi), not area-weighted,',
    '         when strand sizes or prestress levels differ.',
])
sp()
data_table(
    ['x (ft)', 'y_harped (in)', 'y_cg (in)', 'e_cg (in)'],
    [
        ['0 ft', '20.362  (= yc)', '(6+20.362)/2 = 13.181', ('14.362 - 7.181 = 7.181', {'bold': True})],
        ['10 ft', '13.181', '(6+13.181)/2 = 9.590', ('10.772', {'bold': True})],
        ['20 ft  (drape pt)', '6.000', '(6+6)/2 = 6.000', ('14.362', {'bold': True})],
        ['32 ft  (midspan)', '6.000', '(6+6)/2 = 6.000', ('14.362', {'bold': True})],
    ],
    [1.2, 2.0, 2.2, 1.5]
)

heading2('2.3  Allowable Stresses  (CEE 530)')
data_table(
    ['Stage', 'Fiber', 'Limit Type', 'Formula', 'Project 1 Value'],
    [
        ['Transfer', 'Top / Bottom', 'Compression', "+0.60 f'ci", ('+2.880 ksi', {'bold': True})],
        ['Transfer', 'Top / Bottom', 'Tension', "-3*sqrt(f'ci psi)/1000", ('-0.208 ksi', {'bold': True})],
        ['Service (total)', 'Top', 'Compression', "+0.60 f'c", ('+3.600 ksi', {'bold': True})],
        ['Service (Class U)', 'Bottom', 'Tension', "-6*sqrt(f'c psi)/1000", ('-0.465 ksi', {'bold': True})],
    ],
    [1.3, 1.1, 1.1, 1.6, 1.8]
)
sp()
warn([
    [('At Transfer: ', {'bold': True}), 'Use FULL force F (no losses). eta = 1.0 at transfer.'],
    [('At Service: ', {'bold': True}), 'Use EFFECTIVE force Fe = eta * F  where eta = 1 - losses = 0.85 for Project 1.'],
])

heading2('2.4  Fiber Stress Equations')
formula([
    '  AT TRANSFER  (eta = 1.0,  F = Fi,  M = Mi = Msw):',
    '',
    '    f_top = +F/Ac  -  F*e/St  +  Mi/St      (top fiber)',
    '    f_bot = +F/Ac  +  F*e/Sb  -  Mi/Sb      (bottom fiber)',
    '',
    '  AT SERVICE  (eta = 0.85,  Fe = eta*F,  M = MT = Msw+MSDL+MLL):',
    '',
    '    f_top = +Fe/Ac  -  Fe*e/St  +  MT/St    (top fiber)',
    '    f_bot = +Fe/Ac  +  Fe*e/Sb  -  MT/Sb    (bottom fiber)',
    '',
    '  Term-by-term meaning:',
    '    +F/Ac        = uniform compression from prestress (always +)',
    '    -F*e/St      = top fiber TENSION from prestress moment (eccentric force)',
    '    +F*e/Sb      = bottom fiber COMPRESSION from prestress moment',
    '    +M/St        = top fiber COMPRESSION from applied moment',
    '    -M/Sb        = bottom fiber TENSION from applied moment',
])

heading2('2.5  Step-by-Step Stress Check Procedure')
steps = [
    'Compute section properties (or look up): Ac, Ic, yc, St, Sb',
    'Compute e_cg(x) from tendon profile using force-weighted centroid formula',
    'Compute moments: Mi(x) = w_sw*x*(L-x)/2  and  MT(x) = w_s*x*(L-x)/2',
    'Apply fiber stress equations at TRANSFER (eta=1, M=Mi) and SERVICE (eta=0.85, M=MT)',
    'Compare f_top and f_bot against allowable limits. State PASS or FAIL.',
]
for s in steps:
    p = doc.add_paragraph(style='List Number')
    add_run(p, s, size=10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
sp()
warn([
    [('Exam tip - Critical sections to check:', {'bold': True})],
    '  x = 0 (support): M = 0, full eccentric force -> check transfer top tension and bottom compression',
    '  x = L/2 (midspan): maximum moment -> check service bottom tension (usually governs)',
    '  x = drape point (20 ft): transfer bottom compression may govern (large e, small M)',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 3 — ULTIMATE STRENGTH
# ══════════════════════════════════════════════════════════════════════════
heading1('3   Ultimate Strength  (Mn)')

heading2("3.1  beta_1  -  Compression Block Factor  (ACI 318-19 Table 22.2.2.4.3)")
formula([
    "  beta_1 = 0.85                                          f'c <= 4,000 psi",
    "  beta_1 = 0.85 - 0.05*(f'c - 4000)/1000  >= 0.65      4,000 < f'c <= 8,000 psi",
    "  beta_1 = 0.65                                          f'c > 8,000 psi",
    '',
    "  Project 1:  f'c = 6,000 psi",
    "    beta_1 = 0.85 - 0.05*(6000-4000)/1000 = 0.85 - 0.10 = 0.75",
])

heading2('3.2  gamma_p  -  Bonded Tendon Factor  (ACI 318-19 Table 22.3.2.1)')
data_table(
    ['fpy/fpu', 'gamma_p  (ACI 318-19)'],
    [
        ['>= 0.9', '0.28'],
        ['>= 0.8  and  < 0.9', '0.40'],
        ['< 0.8', '0.55'],
    ],
    [3.0, 3.9]
)
p = body('Project 1:  fpy/fpu = 243/270 = 0.90  -->  ')
add_run(p, 'gamma_p = 0.28', bold=True)

heading2('3.3  Prerequisites Before Computing fps')
for txt in [
    "Check fse >= 0.5*fpu:  fse = eta * fpi = 0.85 * 163.4 = 138.9 ksi  >=  0.5*270 = 135 ksi  [OK]",
    'dp = depth to tendon group from COMPRESSION face (top for positive moment)',
    '     Project 1 midspan: dp = h - (yc - e_cg) = 28 - (20.362 - 14.362) = 28 - 6.000 = 22.000 in',
]:
    p = doc.add_paragraph(style='List Bullet')
    add_run(p, txt, size=10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
sp()
warn([
    [('STOP: ', {'bold': True}),
     "dp is measured from the COMPRESSION face (TOP for simply-supported beam under positive moment), NOT from the bottom."],
    ['    dp = h - y_c + e_cg  =  section height - (centroid height) + eccentricity'],
    ['    = 28 - 20.362 + 14.362 = 22.000 in'],
])

heading2('3.4  Stress in Prestressing Steel at Nominal Strength')
body("ACI 318-19 Eq. 22.3.2.1 (bonded tendons, no mild steel in tension zone):")
sp()
formula([
    '  fps = fpu * [1 - (gamma_p/beta_1) * rho_p * fpu/f\'c]',
    '',
    '  where:  rho_p = Aps / (b_eff * dp)',
    "          b_eff = effective flange width for T-section (= 120 in for Project 1)",
    '          dp    = depth to tendon group from compression face (in)',
    '',
    '  General form with mild steel:',
    "  fps = fpu*[1-(gamma_p/beta_1)*(rho_p*fpu/f'c + (d/dp)*omega - (d'/dp)*omega')]",
    "  where omega = rho*fy/f'c (tension), omega' = rho'*fy/f'c (compression)",
])

heading2("3.5  Depth of Equivalent Stress Block  a")
formula([
    '  Rectangular section assumption (check a <= hf first):',
    "  a = Aps * fps / (0.85 * f'c * b_eff)",
    '',
    '  c = a / beta_1              (neutral axis depth from top)',
    '',
    '  T-section check:',
    '  If a <= hf (flange thickness): rectangular formula is valid [OK]',
    '  If a > hf: use T-section formula (see below)',
    '',
    '  T-section formula (when a > hf):',
    "  Ff = 0.85*f'c*(b_eff-b_w)*hf          (flange overhang compression force)",
    "  Aps_w = Aps - Ff/fps                   (equivalent Aps for web only)",
    "  a_w = Aps_w*fps / (0.85*f'c*b_w)      (stress block depth in web)",
    '  Mn = Aps_w*fps*(dp-a_w/2) + Ff*(dp-hf/2)',
])

heading2('3.6  Ductility Check  (epsilon_t)')
formula([
    '  epsilon_t = [(dp - c) / c] * 0.003     (net tensile strain at nominal strength)',
    '',
    '  epsilon_t >= 0.005:  phi = 0.90  (tension-controlled, preferred)',
    '  epsilon_t = 0.004:   phi = 0.81  (ACI 9.3.3, minimum allowed for flexure)',
    '  epsilon_t < 0.004:   Section is compression-controlled -- redesign required',
])

heading2('3.7  Factored Moment  Mu  and Strength Check')
formula([
    '  Mu = 1.2*MD + 1.6*ML               (ACI 318-19 Eq. 5.3.1b)',
    '',
    '  phi*Mn >= Mu                         [Strength requirement]',
    "  phi*Mn >= 1.2*Mcr                    [Minimum reinforcement  ACI 318-19 Sec 24.3.2]",
    '',
    '  Modulus of rupture and cracking moment:',
    "  fr = 7.5*sqrt(f'c psi) / 1000       (ksi, ACI 318-19 Sec 19.2.3)",
    '  Mcr = fr * Sb                        (bottom fiber cracks first under positive moment)',
])

heading2('3.8  Project 1 -- Ultimate Check Summary (midspan, 4 strands)')
two_col([
    ('Total Aps', '4 x 0.153 = 0.612 in^2'),
    ("dp  (from top, = h - yc + e_cg)", '28 - 20.362 + 14.362 = 22.000 in'),
    ("rho_p  = Aps/(b_eff*dp)", '0.612 / (120 x 22.000) = 0.000232'),
    ('gamma_p', '0.28  (fpy/fpu = 0.90 >= 0.9)'),
    ("beta_1  (f'c = 6.0 ksi)", '0.75'),
    ("fps", "270*[1-(0.28/0.75)*0.000232*270/6.0] = 270*(1-0.00462) = 268.75 ksi"),
    ("a  (check vs hf = 2 in)", "0.612*268.75/(0.85*6.0*120) = 0.255 in  < 2 in  [Rectangular OK]"),
    ("c = a/beta_1", "0.255/0.75 = 0.340 in"),
    ('Mn', '0.612*268.75*(22.000-0.255/2) = 0.612*268.75*21.873 = 3,596 kip-in'),
    ('epsilon_t', '[(22.000-0.340)/0.340]*0.003 = 191.6*0.003 = 0.575  >> 0.005  [Tension-controlled]'),
    ('phi', '0.90  (tension-controlled)'),
    ('phi*Mn', '0.90 x 3,596 = 3,237 kip-in  =  269.8 kip-ft'),
    ("fr = 7.5*sqrt(6000)/1000", '0.581 ksi'),
    ('Mcr = fr * Sb', '0.581 x 1701.1 = 988 kip-in'),
    ('1.2*Mcr', '1.2 x 988 = 1,186 kip-in'),
    ('phi*Mn vs 1.2*Mcr', '3,237 >> 1,186  [Min. reinforcement OK]'),
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  SECTION 4 — FEASIBILITY / MAGNEL DIAGRAM
# ══════════════════════════════════════════════════════════════════════════
heading1('4   Feasibility Design Chart  (Magnel Diagram)')

heading2('4.1  Purpose')
body('The Magnel diagram shows all combinations of (1/F, e) that satisfy all four stress constraints simultaneously. '
     'The feasible zone is the intersection of all four regions. '
     'F is the TOTAL prestress force at the section, e = e_cg(x) is the force-weighted eccentricity.')

heading2('4.2  Moments at Each Analysis Section')
formula([
    '  Mi(x) = w_sw * x * (L-x) / 2          [transfer: self-weight only]',
    '  MT(x) = (w_sw + w_SDL + w_LL) * x * (L-x) / 2  [service: all loads]',
])

heading2('4.3  The Four Magnel Lines  (Naaman Numbering)')
body('Starting from the stress inequalities and rearranging to isolate 1/F:')
sp()
formula([
    '  Define denominators (all positive for typical beams):',
    '    dI   = fti + Mi/St        fti = |tension allowable| at transfer (+0.208 ksi)',
    '    dII  = fci + Mi/Sb        fci = compression allowable at transfer (+2.880 ksi)',
    '    dIII = MT/Sb - fts        fts = |tension allowable| at service  (+0.465 ksi)',
    '    dIV  = fcs - MT/St        fcs = compression allowable at service (+3.600 ksi)',
    '',
    '  NOTE: dIII must be > 0 for upper bound to exist. If MT = 0 (at support), dIII <= 0',
    '         and Line III gives no constraint -> F_min = 0 at that section.',
])
sp()
data_table(
    ['Line', 'Condition', 'Inequality', 'Bound type', 'Governs'],
    [
        ['I', 'Top @ Transfer, Tension >= -fti',
         '1/F <= (e/St - 1/Ac) / dI  [if dI>0, upper bound]',
         'UPPER bound on 1/F', 'Near support'],
        ['II', 'Bot @ Transfer, Compr. <= +fci',
         '1/F >= (1/Ac + e/Sb) / dII',
         'LOWER bound on 1/F', 'Drape point'],
        ['III', 'Bot @ Service, Tension >= -fts',
         '1/F <= eta*(1/Ac + e/Sb) / dIII',
         ('UPPER bound on 1/F', {'bold': True}), ('Midspan (governs F_min)', {'bold': True})],
        ['IV', 'Top @ Service, Compr. <= +fcs',
         '1/F >= eta*(1/Ac - e/St) / dIV',
         'LOWER bound on 1/F', 'Midspan'],
    ],
    [0.4, 1.9, 2.3, 1.5, 1.8]
)
sp()
warn([
    [('Feasible zone: ', {'bold': True}),
     'max(Line II, Line IV) <= 1/F <= min(Line I, Line III)'],
    ['  AND  1/F > 0   AND  e_geo_min <= e <= e_geo_max  where  e_geo_max = yb - cover'],
    [('PITFALL: ', {'bold': True}),
     'If denominator dI or dIV goes negative as e varies, the inequality FLIPS. Always check sign.'],
])

heading2('4.4  Finding F_min at a Given Section and Eccentricity')
formula([
    '  Line III sets the UPPER bound on 1/F.  Larger 1/F = smaller F.',
    '    (1/F)_max = eta * (1/Ac + e/Sb) / dIII',
    '',
    '  Minimum required prestress force:',
    '    F_min = 1 / (1/F)_max = dIII / [eta * (1/Ac + e/Sb)]',
    '',
    '  Maximum allowed prestress force (from Line II):',
    '    F_max = dII / (1/Ac + e/Sb)',
])

heading2('4.5  Governing Section and Number of Strands')
steps2 = [
    'Compute Mi(x) and MT(x) at all analysis sections',
    'Compute e_cg(x) at each section from the tendon profile',
    'Compute dIII at each section.  If dIII <= 0: F_min = 0 at that section',
    'Compute F_min at each section.  The LARGEST F_min governs (usually midspan)',
    'Number of strands: n = ceil(F_min / Pi) where Pi = Aps * fpi per strand',
    'Check F_provided = n * Pi against F_max at ALL sections -- must not exceed upper bound anywhere',
]
for s in steps2:
    p = doc.add_paragraph(style='List Number')
    add_run(p, s, size=10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

heading2('4.6  Project 1 -- Feasibility Results')
data_table(
    ['Section  x', 'Mi (kip-in)', 'MT (kip-in)', 'e_cg (in)', 'F_min (kip)', 'F_max (kip)'],
    [
        ['0 ft', '0', '0', '7.181', ('0', {'bold': True}), ('459', {'bold': True})],
        ['10 ft', '1,643.6', '3,814.4', '10.772', ('249', {'bold': True}), ('459', {'bold': True})],
        ['20 ft  (drape)', '2,678.5', '6,216.1', '14.362', ('357', {'bold': True}), ('424.4', {'bold': True})],
        ['32 ft  (midspan)', '3,116.8', '7,233.3', '14.362', ('424.5', {'bold': True}), ('449', {'bold': True})],
    ],
    [1.1, 1.0, 1.1, 0.9, 1.0, 1.0]
)
sp()
formula([
    '  Governing:  x = 32 ft  (midspan)',
    '    F_min = 424.5 kip',
    '    n = ceil(424.5 / 25.0) = ceil(16.98) = 17 strands',
    '    F_provided = 17 x 25.0 = 425 kip',
    '',
    '  Note on F_max conflict:',
    '    F_max at x=20 ft = 424.4 kip  <  F_min at x=32 ft = 424.5 kip',
    '    Gap = 0.1 kip (0.2% overage). For engineering practice: 17 strands is the only choice.',
])
sp()
warn([
    [('KEY PITFALLS in Feasibility Problems:', {'bold': True})],
    '  1. If dIII <= 0: no service tension upper bound -> F_min = 0 at that section (x=0, MT=0)',
    '  2. Line III gives UPPER bound on 1/F = LOWER bound on F. Lower 1/F = higher F = stronger beam.',
    '  3. Line II gives LOWER bound on 1/F = UPPER bound on F. Check F <= F_max at ALL sections.',
    '  4. At transfer use full F (no losses). At service use Fe = eta*F.',
    '  5. e_cg uses FORCE-weighted average (Pi = Aps*fpi), not area average, when strand sizes differ.',
    '  6. Denominator formulas use MAGNITUDE of allowable stresses (treat fti and fts as positive).',
])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
#  APPENDIX
# ══════════════════════════════════════════════════════════════════════════
heading1('Appendix  -  Quick Reference Card  (Project 1 Numbers)')

heading2('A.1  Material Properties')
data_table(
    ['Property', 'Symbol', 'Value', 'Unit'],
    [
        ["Concrete 28-day strength",      "f'c",              "6.0",    "ksi"],
        ["Concrete strength at transfer", "f'ci",             "4.8",    "ksi"],
        ["Strand ultimate strength",      "fpu",              "270",    "ksi"],
        ["Strand yield strength",         "fpy = 0.90*fpu",   "243",    "ksi"],
        ["Strand initial prestress",      "fpi = Pi/Aps",     "163.4",  "ksi"],
        ["Strand area (0.5-in dia.)",     "Aps",              "0.153",  "in^2"],
        ["Initial force per strand",      "Pi = Aps*fpi",     "25.0",   "kip"],
        ["Prestress losses",              "losses",           "15%",    "fraction 0.15"],
        ["Loss factor",                   "eta = 1-losses",   "0.85",   "dimensionless"],
        ["Effective force/strand",        "Pe = eta*Pi",      "21.25",  "kip"],
        ["Number of strands",             "n  (design)",      "17",     "strands"],
        ["Total initial force",           "F = n*Pi",         "425",    "kip"],
        ["Total effective force",         "Fe = eta*F",       "361.25", "kip"],
    ],
    [2.2, 1.7, 1.0, 1.2]
)

heading2('A.2  Allowable Stresses  (CEE 530 Class U)')
data_table(
    ['Condition', 'Limit Formula', 'Project 1 Value'],
    [
        ['Transfer -- Compression', "+0.60 * f'ci",             '+2.880 ksi'],
        ['Transfer -- Tension',     "-3*sqrt(f'ci psi)/1000",   '-0.208 ksi'],
        ['Service Total -- Compr.', "+0.60 * f'c",              '+3.600 ksi'],
        ['Service -- Tension (U)',  "-6*sqrt(f'c psi)/1000",    '-0.465 ksi'],
    ],
    [2.1, 2.3, 1.7]
)

heading2('A.3  Section Properties  (Double-T)')
data_table(
    ['Property', 'Value', 'Unit'],
    [
        ['Ac', '487.0', 'in^2'],
        ['Ic', '34,638.8', 'in^4'],
        ['yc  (centroid from bottom)', '20.362', 'in'],
        ['yt  (centroid to top fiber)', '7.638', 'in'],
        ['yb  (= yc, centroid to bot)', '20.362', 'in'],
        ['St  (top section modulus)', '4,535.1', 'in^3'],
        ['Sb  (bottom section modulus)', '1,701.1', 'in^3'],
        ['kt  (kern toward top)', '9.312', 'in'],
        ['kb  (kern toward bottom)', '3.493', 'in'],
        ['r^2 = Ic/Ac', '71.13', 'in^2'],
    ],
    [2.5, 2.0, 1.6]
)

heading2('A.4  Tendon Layout  (Project 1, 4 strands at midspan)')
data_table(
    ['Tendon', 'Type', 'y at support', 'y at drape (x=240in)', 'e at midspan'],
    [
        ['1, 3 (outer)', 'Straight', '6.00 in', '6.00 in', '14.362 in'],
        ['2, 4 (inner)', 'Harped', '20.362 in (= yc)', '6.00 in', '14.362 in'],
    ],
    [0.9, 0.9, 1.0, 1.5, 1.0]
)

# ── Save ───────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'Saved: {OUT}')
